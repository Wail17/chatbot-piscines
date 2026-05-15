"""
Build Word documents (.docx) from the project's markdown documentation.

Produces 4 styled Word documents ready for OneDrive upload:
- Wifipool_Chatbot_README.docx
- Wifipool_Chatbot_Architecture.docx
- Wifipool_Chatbot_Security.docx
- Wifipool_Chatbot_API.docx

Plus a combined version: Wifipool_Chatbot_Documentation.docx

Run: python build_docs.py
"""
from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent

DOCS: List[Tuple[str, str, str]] = [
    ("Wifipool_Chatbot_README.docx",        "README.md",                 "Wifipool AI Assistant — Project Overview"),
    ("Wifipool_Chatbot_Architecture.docx",  "docs/ARCHITECTURE.md",      "System Architecture"),
    ("Wifipool_Chatbot_Security.docx",      "docs/SECURITY.md",          "Security Audit"),
    ("Wifipool_Chatbot_API.docx",           "docs/API.md",               "API Reference"),
]

COMBINED_NAME = "Wifipool_Chatbot_Documentation.docx"

OCEAN_500 = RGBColor(0x0E, 0x7A, 0xC4)
OCEAN_700 = RGBColor(0x07, 0x4B, 0x80)
INK_900   = RGBColor(0x0B, 0x25, 0x40)
INK_500   = RGBColor(0x62, 0x7A, 0x95)
CODE_BG   = "F1F6FB"
TABLE_HDR = "0E7AC4"
TABLE_HDR_INK = RGBColor(0xFF, 0xFF, 0xFF)


# ─────────────────────────────────────────────────────────────────────────────
# Style helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_paragraph_bg(paragraph, hex_color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def _apply_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "B9CEE0")
        borders.append(b)
    tbl_pr.append(borders)


def _configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK_900
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for level, size, color in [
        ("Heading 1", 22, OCEAN_700),
        ("Heading 2", 17, OCEAN_700),
        ("Heading 3", 14, OCEAN_500),
        ("Heading 4", 12, OCEAN_500),
    ]:
        st = styles[level]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(6)

    # Custom code style
    if "CodeBlock" not in [s.name for s in styles]:
        code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code.base_style = styles["Normal"]
        code.font.name = "Consolas"
        code.font.size = Pt(9.5)
        code.font.color.rgb = RGBColor(0x1C, 0x4A, 0x6E)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.left_indent = Cm(0.4)
        code.paragraph_format.line_spacing = 1.15

    if "InlineCode" not in [s.name for s in styles]:
        ic = styles.add_style("InlineCode", WD_STYLE_TYPE.CHARACTER)
        ic.font.name = "Consolas"
        ic.font.size = Pt(10)
        ic.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)


# ─────────────────────────────────────────────────────────────────────────────
# Inline formatting
# ─────────────────────────────────────────────────────────────────────────────

INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+?\*\*|__[^_]+?__|`[^`]+?`|\[[^\]]+?\]\([^)]+?\)|\*[^*]+?\*|_[^_]+?_)"
)


def _add_inline_runs(paragraph, text: str):
    text = text.replace("✅", "✓").replace("❌", "✗").replace("🟡", "●").replace("🟢", "●")
    text = re.sub(r"[\U0001F300-\U0001FAFF\U00002600-\U000027BF]", "", text)
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif part.startswith("[") and "](" in part:
            label, _, rest = part[1:].partition("](")
            url = rest.rstrip(")")
            run = paragraph.add_run(label)
            run.font.color.rgb = OCEAN_500
            run.underline = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("_") and part.endswith("_") and not part.startswith("__"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ─────────────────────────────────────────────────────────────────────────────
# Markdown parser → Word
# ─────────────────────────────────────────────────────────────────────────────

def _parse_table(lines: List[str]) -> List[List[str]]:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(set(c) <= set("-: ") for c in cells):
            continue
        rows.append(cells)
    return rows


def _is_table_separator(line: str) -> bool:
    s = line.strip().strip("|")
    return bool(s) and all(set(c.strip()) <= set("-: ") for c in s.split("|"))


def _render_markdown(doc: Document, md_text: str):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            # Mermaid → render as labeled box
            if lang.lower() == "mermaid":
                p = doc.add_paragraph()
                p.add_run("[ Diagram — see GitHub for live rendering ]").italic = True
                p.runs[0].font.color.rgb = INK_500
                p.runs[0].font.size = Pt(9)
            for code_line in buf:
                p = doc.add_paragraph(style="CodeBlock")
                p.add_run(code_line or " ")
                _set_paragraph_bg(p, CODE_BG)
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                text = re.sub(r"^[^\w]+", "", text).strip()
                heading = doc.add_heading(text, level=min(level, 4))
                heading.paragraph_format.keep_with_next = True
                i += 1
                continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", stripped):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            p_pr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "B9CEE0")
            pbdr.append(bottom)
            p_pr.append(pbdr)
            i += 1
            continue

        # Block quote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            _set_paragraph_bg(p, "EEF8FD")
            text = stripped.lstrip(">").strip()
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = OCEAN_700
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            rows = _parse_table(table_lines)
            if rows:
                max_cols = max(len(r) for r in rows)
                rows = [r + [""] * (max_cols - len(r)) for r in rows]
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                _apply_borders(table)
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        if r_idx == 0:
                            _set_cell_bg(cell, TABLE_HDR)
                            run = p.add_run(cell_text)
                            run.bold = True
                            run.font.color.rgb = TABLE_HDR_INK
                            run.font.size = Pt(10.5)
                        else:
                            _add_inline_runs(p, cell_text)
                            for run in p.runs:
                                run.font.size = Pt(10)
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
            i = j
            continue

        # Bullet / numbered list
        if re.match(r"^\s*[-*+]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            content = re.sub(r"^\s*[-*+]\s+", "", line)
            _add_inline_runs(p, content)
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            _add_inline_runs(p, content)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, line)
        i += 1


# ─────────────────────────────────────────────────────────────────────────────
# Document assembly
# ─────────────────────────────────────────────────────────────────────────────

def _title_page(doc: Document, title: str, subtitle: str):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WIFIPOOL")
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = OCEAN_700

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI Assistant")
    run.font.size = Pt(22)
    run.font.color.rgb = OCEAN_500

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = INK_900

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = INK_500

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Internship Deliverable — Backend Engineering & Security")
    run.font.size = Pt(11)
    run.font.color.rgb = INK_500

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Wifipool / Beniferro Belgium · 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = INK_500

    doc.add_page_break()


def build_one(md_path: Path, out_path: Path, title: str, subtitle: str):
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    _configure_styles(doc)
    _title_page(doc, title, subtitle)
    _render_markdown(doc, text)
    doc.save(out_path)
    print(f"  [ok] {out_path.name}  ({len(text):,} chars md -> {out_path.stat().st_size // 1024} KB docx)")


def build_combined(out_path: Path):
    doc = Document()
    _configure_styles(doc)
    _title_page(doc, "Complete Technical Documentation", "Project overview · Architecture · Security · API")

    for idx, (_, md_rel, title) in enumerate(DOCS):
        md_path = ROOT / md_rel
        if not md_path.exists():
            continue
        # Section header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"PART {idx + 1}")
        run.font.size = Pt(14)
        run.font.color.rgb = INK_500
        run.font.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(24)
        run.bold = True
        run.font.color.rgb = OCEAN_700
        doc.add_paragraph()

        _render_markdown(doc, md_path.read_text(encoding="utf-8"))
        if idx < len(DOCS) - 1:
            doc.add_page_break()

    doc.save(out_path)
    print(f"  [ok]{out_path.name}  ({out_path.stat().st_size // 1024} KB combined)")


def main():
    out_dir = ROOT / "deliverables"
    out_dir.mkdir(exist_ok=True)
    print(f"Building Word documents into {out_dir}/")
    print("-" * 60)
    for out_name, md_rel, title in DOCS:
        md_path = ROOT / md_rel
        if not md_path.exists():
            print(f"  [!!] {md_rel} not found - skipping")
            continue
        out_path = out_dir / out_name
        subtitle = {
            "README.md":                  "Wifipool AI Assistant — Project Overview",
            "docs/ARCHITECTURE.md":       "System Architecture",
            "docs/SECURITY.md":           "Security Audit",
            "docs/API.md":                "API Reference",
        }[md_rel]
        build_one(md_path, out_path, title, subtitle)
    print("-" * 60)
    print("Building combined edition...")
    build_combined(out_dir / COMBINED_NAME)
    print("-" * 60)
    print("Done. Ready to upload to OneDrive.")


if __name__ == "__main__":
    main()
